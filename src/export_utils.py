"""
Utility functions for exporting sponsor data in various formats
"""
import io
import json
import re
import pandas as pd

def create_safe_filename(title):
    """Create a safe filename from video title."""
    safe_title = re.sub(r'[\\/*?:"<>|]', "", title)
    return safe_title[:50].strip()

def create_json_export(sponsor_info):
    """Create JSON export data."""
    return json.dumps(sponsor_info, indent=2)

def create_csv_export(sponsor_info):
    """Create CSV export data."""
    df = pd.DataFrame(sponsor_info)
    csv_buffer = io.StringIO()
    df.to_csv(csv_buffer, index=False)
    return csv_buffer.getvalue()

def create_excel_export(sponsor_info):
    """Create Excel export data."""
    df = pd.DataFrame(sponsor_info)
    excel_buffer = io.BytesIO()
    df.to_excel(excel_buffer, index=False, engine="openpyxl")
    excel_buffer.seek(0)
    return excel_buffer

def create_text_export(sponsor_info):
    """Create text export data."""
    text_content = "YouTube Sponsor Information\n"
    text_content += "=" * 30 + "\n\n"
    for item in sponsor_info:
        text_content += f"Brand: {item['Brand']}\n"
        text_content += f"URL: {item['URL']}\n"
        text_content += "-" * 30 + "\n\n"
    return text_content

def create_docx_export(sponsor_info):
    """Create DOCX export data."""
    try:
        from docx import Document
        
        # Create document
        doc = Document()
        doc.add_heading("YouTube Sponsor Information", 0)
        
        # Add table
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Brand"
        header_cells[1].text = "URL"
        
        # Add data rows
        for item in sponsor_info:
            row_cells = table.add_row().cells
            row_cells[0].text = item["Brand"]
            row_cells[1].text = item["URL"]
        
        # Save to buffer
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_buffer.seek(0)
        
        return docx_buffer
    except ImportError:
        raise ImportError("DOCX export requires python-docx library")
