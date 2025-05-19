"""
YouTube Sponsor Extractor - A streamlined app to extract sponsor information from YouTube video descriptions
"""
import os
import re
import json
import io
import streamlit as st
import pandas as pd
from googleapiclient.discovery import build
from googleapiclient.errors import HttpError
from openai import OpenAI
from dotenv import load_dotenv

# Load environment variables
load_dotenv()

# Constants
PROMPT_TEMPLATE = """
You are a specialized extraction tool that identifies sponsor information from YouTube video descriptions.

Your task:
1. Extract all sponsor information from the provided YouTube video description
2. Look for any sponsored products, affiliate links, or promotional content
3. For each sponsor, extract the brand name and URL
4. Return ONLY a valid JSON array of objects with 'Brand' and 'URL' keys
5. If no sponsors are found, return an empty array []
6. Do not include any explanations, markdown formatting, or text outside the JSON

Description:
{description}

Respond with ONLY the JSON array:
"""

# YouTube API functions
def extract_video_id(url):
    """Extract video ID from YouTube URL."""
    patterns = [r"(?:v=|\/)([0-9A-Za-z_-]{11}).*", r"youtu\.be\/([0-9A-Za-z_-]{11})"]
    for pattern in patterns:
        match = re.search(pattern, url)
        if match:
            return match.group(1)
    return None

def fetch_youtube_description(youtube_url):
    """Fetch the description of a YouTube video using the YouTube Data API."""
    api_key = os.getenv("YOUTUBE_API_KEY")
    if not api_key:
        raise ValueError("YouTube API key not found in .env file")

    video_id = extract_video_id(youtube_url)
    if not video_id:
        raise ValueError("Invalid YouTube URL")

    try:
        youtube = build("youtube", "v3", developerKey=api_key)
        request = youtube.videos().list(part="snippet", id=video_id)
        response = request.execute()

        if "items" in response and len(response["items"]) > 0:
            return response["items"][0]["snippet"]["description"]
        return None
    except HttpError as e:
        raise Exception(f"YouTube API error: {str(e)}")

# OpenRouter API function
def extract_sponsor_info(description):
    """Use OpenRouter API to extract sponsor information from the description."""
    api_key = os.getenv("OPENROUTER_API_KEY")
    if not api_key:
        raise ValueError("OpenRouter API key not found in .env file")

    prompt = PROMPT_TEMPLATE.replace("{description}", description)

    client = OpenAI(
        base_url="https://openrouter.ai/api/v1",
        api_key=api_key,
    )

    try:
        completion = client.chat.completions.create(
            model="google/gemma-3-27b-it:free",
            messages=[{"role": "user", "content": prompt}],
            response_format={"type": "json_object"},  # Ensure JSON response format
            extra_headers={
                "HTTP-Referer": "http://localhost:8501",
                "X-Title": "YouTube Sponsor Extractor",
            },
        )
        
        # Get the response content
        response_content = completion.choices[0].message.content
        
        # Debug the response content
        if not response_content or response_content.isspace():
            return []  # Return empty array if response is empty
            
        # Try to parse the JSON, with fallback handling
        try:
            # Clean the response if it has markdown code blocks
            if response_content.strip().startswith('```json') and response_content.strip().endswith('```'):
                # Extract content from markdown code block
                response_content = response_content.strip().removeprefix('```json').removesuffix('```').strip()
            elif response_content.strip().startswith('```') and response_content.strip().endswith('```'):
                # Extract content from generic code block
                response_content = response_content.strip().removeprefix('```').removesuffix('```').strip()
                
            sponsor_info = json.loads(response_content)
            
            # Validate the response structure and clean the data
            cleaned_sponsors = []
            
            if isinstance(sponsor_info, list):
                raw_sponsors = sponsor_info
            elif isinstance(sponsor_info, dict) and 'sponsors' in sponsor_info:
                raw_sponsors = sponsor_info['sponsors']
            else:
                # If it's not in the expected format, return an empty list
                return []
                
            # Clean the data - remove entries with <NA> or None values
            for sponsor in raw_sponsors:
                if not isinstance(sponsor, dict):
                    continue
                    
                # Ensure both Brand and URL keys exist
                if 'Brand' not in sponsor or 'URL' not in sponsor:
                    continue
                    
                # Skip entries with empty, None, or <NA> values
                brand = sponsor.get('Brand')
                url = sponsor.get('URL')
                
                if brand in (None, '', '<NA>', 'NA', 'N/A') or url in (None, '', '<NA>', 'NA', 'N/A'):
                    continue
                    
                # Add clean entry
                cleaned_sponsors.append({
                    'Brand': brand,
                    'URL': url if url else ''
                })
                
            return cleaned_sponsors
                
        except json.JSONDecodeError:
            # If JSON parsing fails, return empty list
            return []
        except Exception as e:
            # Handle any other errors in processing
            print(f"Error processing sponsor data: {str(e)}")
            return []
            
    except Exception as e:
        raise Exception(f"OpenRouter API error: {str(e)}")

# Main Streamlit app

# Main Streamlit app
def main():
    # Set page config for tab title
    st.set_page_config(
        page_title="YouTube Sponsor Extractor",
        page_icon="🎬",
        layout="centered"
    )
    
    st.title("YouTube Sponsor Extractor")
    st.write(
        "Extract sponsor information (brand name and URL) from YouTube video descriptions."
    )

    # Input YouTube URL
    youtube_url = st.text_input(
        "YouTube URL", placeholder="https://www.youtube.com/watch?v=VIDEO_ID"
    )

    # Process button
    if st.button("Extract Sponsors"):
        if not youtube_url:
            st.write("Please enter a YouTube URL.")
            return
            
        # Validate URL
        if not re.match(r"https?://(www\.)?(youtube\.com|youtu\.be)/", youtube_url):
            st.error("Please enter a valid YouTube URL.")
            return

        try:
            # Fetch video information
            with st.spinner("Fetching video information..."):
                # Get video data
                youtube = build("youtube", "v3", developerKey=os.getenv("YOUTUBE_API_KEY"))
                video_id = extract_video_id(youtube_url)
                request = youtube.videos().list(part="snippet", id=video_id)
                response = request.execute()
                
                if "items" in response and len(response["items"]) > 0:
                    snippet = response["items"][0]["snippet"]
                    description = snippet["description"]
                    video_title = snippet["title"]
                else:
                    st.error("Could not fetch video information. Check the URL or API key.")
                    return
            
            # Extract sponsor info
            with st.spinner("Extracting sponsor information..."):
                sponsor_info = extract_sponsor_info(description)

            # Display results
            if sponsor_info:
                st.write(f"### Video: {video_title}")
                st.write("### Extracted Sponsor Information")
                df = pd.DataFrame(sponsor_info)
                st.table(df)

                # Prepare download options
                st.write("### Download Results")
                
                # Create columns for download buttons
                col1, col2, col3, col4, col5 = st.columns(5)
                
                # Create sanitized filename from video title
                safe_title = re.sub(r'[\\/*?:"<>|]', "", video_title)
                safe_title = safe_title[:50].strip()
                
                # JSON download
                with col1:
                    json_data = json.dumps(sponsor_info, indent=2)
                    st.download_button(
                        label="JSON",
                        data=json_data,
                        file_name=f"{safe_title}_sponsors.json",
                        mime="application/json"
                    )
                
                # CSV download
                with col2:
                    csv_buffer = io.StringIO()
                    df.to_csv(csv_buffer, index=False)
                    st.download_button(
                        label="CSV",
                        data=csv_buffer.getvalue(),
                        file_name=f"{safe_title}_sponsors.csv",
                        mime="text/csv"
                    )
                
                # Excel download
                with col3:
                    excel_buffer = io.BytesIO()
                    df.to_excel(excel_buffer, index=False, engine="openpyxl")
                    excel_buffer.seek(0)
                    st.download_button(
                        label="Excel",
                        data=excel_buffer,
                        file_name=f"{safe_title}_sponsors.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
                    )
                
                # TXT download
                with col4:
                    text_content = "YouTube Sponsor Information\n"
                    text_content += "=" * 30 + "\n\n"
                    for item in sponsor_info:
                        text_content += f"Brand: {item['Brand']}\n"
                        text_content += f"URL: {item['URL']}\n"
                        text_content += "-" * 30 + "\n\n"
                    st.download_button(
                        label="TXT",
                        data=text_content,
                        file_name=f"{safe_title}_sponsors.txt",
                        mime="text/plain"
                    )
                
                # DOCX download
                with col5:
                    try:
                        from docx import Document
                        
                        # Create document
                        doc = Document()
                        doc.add_heading("YouTube Sponsor Information", 0)
                        
                        # Add table
                        table = doc.add_table(rows=1, cols=2)
                        table.style = 'Table Grid'
                        
                        # Add header row
                        header_cells = table.rows[0].cells
                        header_cells[0].text = "Brand"
                        header_cells[1].text = "URL"
                        
                        # Add data rows
                        for item in sponsor_info:
                            row_cells = table.add_row().cells
                            row_cells[0].text = item["Brand"]
                            row_cells[1].text = item["URL"]
                        
                        # Save to buffer
                        docx_buffer = io.BytesIO()
                        doc.save(docx_buffer)
                        docx_buffer.seek(0)
                        
                        st.download_button(
                            label="DOCX",
                            data=docx_buffer,
                            file_name=f"{safe_title}_sponsors.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except ImportError:
                        st.error("DOCX export requires python-docx library")
            else:
                st.write(f"### Video: {video_title}")
                st.write("No sponsor information found in the description.")
                
        except Exception as e:
            st.error(f"An error occurred: {str(e)}")

if __name__ == "__main__":
    main()
